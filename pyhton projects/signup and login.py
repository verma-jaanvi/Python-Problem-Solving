import tkinter as tk
from tkinter import messagebox
from docx import Document
import os


def validate_username_password(username, password):
    if ' ' in username:
        messagebox.showerror("Input Error", "Username cannot contain spaces.")
        return False
    if len(username) < 3:
        messagebox.showerror("Input Error", "Username must be at least 3 characters long.")
        return False
    if not password:
        messagebox.showerror("Input Error", "Password cannot be empty.")
        return False
    if len(password) < 6:
        messagebox.showerror("Input Error", "Password must be at least 6 characters long.")
        return False
    return True
def open_register_window():
    register_window = tk.Toplevel(window)
    register_window.title("Sign Up")
    register_window.geometry("350x300")

    conditions_label = tk.Label(register_window, text="Conditions:\n"
                                                     "- Username: At least 3 characters, no spaces.\n"
                                                     "- Password: At least 6 characters.", 
                                fg="blue", justify="left")
    conditions_label.pack(pady=10)

    def save_signup_data():
        new_username = new_username_entry.get()
        new_password = new_password_entry.get()

        if validate_username_password(new_username, new_password):
            if os.path.exists('registered_users.docx'):
                doc = Document('registered_users.docx')
            else:
                doc = Document()

            doc.add_heading('Registered Users', 0)
            doc.add_paragraph(f'Username: {new_username}')
            doc.add_paragraph(f'Password: {new_password}')
            doc.save('registered_users.docx')
            messagebox.showinfo("Success", "Sign up successful! Data saved in registered_users.docx")

    new_username_label = tk.Label(register_window, text="New Username:")
    new_username_label.pack(pady=5)
    new_username_entry = tk.Entry(register_window)
    new_username_entry.pack(pady=5)

    new_password_label = tk.Label(register_window, text="New Password:")
    new_password_label.pack(pady=5)
    new_password_entry = tk.Entry(register_window, show="*")
    new_password_entry.pack(pady=5)

    signup_button = tk.Button(register_window, text="Sign Up", command=save_signup_data)
    signup_button.pack(pady=20)

def login():
    username = username_entry.get()
    password = password_entry.get()
    if validate_username_password(username, password):
        if os.path.exists('registered_users.docx'):
            doc = Document('registered_users.docx')
            users_data = [para.text for para in doc.paragraphs]
            
            for i in range(0, len(users_data), 2):
                try:
                    saved_username = users_data[i].split(': ')[1]
                    saved_password = users_data[i+1].split(': ')[1]
                    
                    if saved_username == username and saved_password == password:
                        messagebox.showinfo("Success", "Login successful!")
                        return
                except IndexError:
                    messagebox.showerror("Data Error", "Corrupted user data detected. Please register again.")
                    break

            messagebox.showinfo("Info", "User not found. Redirecting to Sign Up.")
            open_register_window()
        else:
            messagebox.showinfo("Info", "No registered users. Redirecting to Sign Up.")
            open_register_window()

window = tk.Tk()
window.title("Login or Sign Up")
window.geometry("350x250")

conditions_label = tk.Label(window, text="Conditions:\n"
                                        "- Username: At least 3 characters, no spaces.\n"
                                        "- Password: At least 6 characters.",
                            fg="blue", justify="left")
conditions_label.pack(pady=10)

username_label = tk.Label(window, text="Username:")
username_label.pack(pady=5)
username_entry = tk.Entry(window)
username_entry.pack(pady=5)

password_label = tk.Label(window, text="Password:")
password_label.pack(pady=5)
password_entry = tk.Entry(window, show="*")
password_entry.pack(pady=5)

login_button = tk.Button(window, text="Login", command=login)
login_button.pack(pady=10)

signup_button = tk.Button(window, text="Sign Up", command=open_register_window)
signup_button.pack(pady=10)

window.mainloop()
