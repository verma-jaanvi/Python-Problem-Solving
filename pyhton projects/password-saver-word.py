import tkinter as tk
from tkinter import messagebox
from docx import Document
import os

def save_user_data(username, password, filename):
    """Saves the username and password to a Word document."""
    doc = Document()
    doc.add_heading('User  Data', 0)
    doc.add_paragraph(f'Username: {username}')
    doc.add_paragraph(f'Password: {password}')
    
    try:
        doc.save(filename)
        messagebox.showinfo("Success", f"Data saved successfully in {filename}")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to save data: {e}")

def handle_submit():
    """Handles the submission of data from the sign-in window."""
    username = username_entry.get()
    password = password_entry.get()

    if username and password:
        save_user_data(username, password, 'user_data.docx')
        window.destroy() 
    else:
        messagebox.showwarning("Input Error", "Please enter both username and password")

def signin():
    """Creates a new window for user to sign in."""
    global username_entry, password_entry
    window = tk.Tk()
    window.title("Sign in")
    window.geometry("300x200")

    tk.Label(window, text="Username:").pack(pady=10)
    username_entry = tk.Entry(window)
    username_entry.pack(pady=5)

    tk.Label(window, text="Password:").pack(pady=10)
    password_entry = tk.Entry(window, show="*")
    password_entry.pack(pady=5)

    tk.Button(window, text="Submit", command=handle_submit).pack(pady=10)
  
def save_data():
    """Handles saving data from the main window."""
    username = username_entry.get()
    password = password_entry.get()

    if username and password:
        save_user_data(username, password, 'new_data.docx')
        window.destroy()  # Close the main window
    else:
        messagebox.showwarning("Input Error", "Please enter both username and password")

# Main window setup
window = tk.Tk()
window.title("Username and Password Saver")
window.geometry("300x280")

tk.Label(window, text="Username:").pack(pady=10)
username_entry = tk.Entry(window)
username_entry.pack(pady=5)

tk.Label(window, text="Password:").pack(pady=10)
password_entry = tk.Entry(window, show="*")
password_entry.pack(pady=5)

tk.Button(window, text="Log in", command=save_data).pack(pady=10)

tk.Label(window, text="Don't have an account?").pack(pady=5)
tk.Button(window, text="Sign in", command=signin).pack(pady=15)

window.mainloop()